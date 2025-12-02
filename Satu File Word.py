import streamlit as st
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO
import re

st.title("Generator Semua Surat PNS DLH (1 File Word)")

# Upload file
excel_file = st.file_uploader("📄 Upload Excel (daftar pegawai)", type="xlsx")
template_file = st.file_uploader("📄 Upload Template Surat (Word .docx)", type="docx")


# Fungsi menggabungkan dokumen
def merge_documents(docs):
    final_doc = Document()

    for index, doc_stream in enumerate(docs):
        sub_doc = Document(doc_stream)

        if index > 0:
            final_doc.add_page_break()

        for element in sub_doc.element.body:
            final_doc.element.body.append(element)

    output = BytesIO()
    final_doc.save(output)
    output.seek(0)
    return output


# Fungsi menghasilkan semua surat dalam 1 file
def generate_single_docx(template_bytes, data_rows):
    rendered_docs = []

    for row in data_rows:
        tpl = DocxTemplate(BytesIO(template_bytes))

        context = {
            "Nama_Pegawai": row.get("Nama_Pegawai", ""),
            "Jabatan": row.get("Jabatan", ""),
            "Golongan": row.get("Golongan", ""),
            "Pangkat": row.get("Pangkat", ""),
            "NIP": row.get("NIP", ""),
        }

        tpl.render(context)

        doc_io = BytesIO()
        tpl.save(doc_io)
        doc_io.seek(0)

        rendered_docs.append(BytesIO(doc_io.read()))

    return merge_documents(rendered_docs)


# Preview surat pertama
def preview_docx_from_template(template_bytes, context):
    tpl = DocxTemplate(BytesIO(template_bytes))
    tpl.render(context)

    temp = BytesIO()
    tpl.save(temp)
    temp.seek(0)

    doc = Document(temp)
    return "\n".join([p.text for p in doc.paragraphs])


# Logika utama aplikasi
if excel_file and template_file:
    df = pd.read_excel(excel_file).fillna("")

    required_columns = ["Nama_Pegawai"]
    missing_cols = [col for col in required_columns if col not in df.columns]

    if missing_cols:
        st.error(f"Kolom berikut wajib ada di Excel: {', '.join(missing_cols)}")
    else:
        df_valid = df[df["Nama_Pegawai"].str.strip() != ""]

        if df_valid.empty:
            st.error("Tidak ada data pegawai yang valid. Kolom 'Nama_Pegawai' wajib diisi.")
        else:
            data_rows = df_valid.to_dict(orient="records")
            template_bytes = template_file.read()

            st.subheader("📄 Preview Isi Surat Pertama:")
            st.text_area(
                "Preview Surat",
                preview_docx_from_template(template_bytes, data_rows[0]),
                height=500
            )

            # Download satu dokumen berisi semua surat
            st.subheader("📥 Download Dokumen Gabungan")
            final_doc = generate_single_docx(template_bytes, data_rows)

            st.download_button(
                label="📥 Download Semua Surat Dalam Satu File (.docx)",
                data=final_doc,
                file_name="Semua_Surat_Pegawai_DLH.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
