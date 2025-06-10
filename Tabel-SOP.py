import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from io import BytesIO

# Konfigurasi Streamlit
st.set_page_config(layout="wide")
st.title("📋 SOP - Penerbitan Izin Pemanfaatan Air Limbah")
st.markdown("### Bagian Proses dan Pelaksana (dengan simbol visual sederhana)")

# Simbol
kotak = "🟦"
diamond = "🔷"

# Data SOP
data = [
    ["1", "Pemohon ajukan permohonan", kotak, "", "", "", "", "", "Dokumen Permohonan", "15 menit", "Tanda Terima"],
    ["2", "Meneruskan surat", "", kotak, "", "", "", "", "Surat", "15 menit", "-"],
    ["3", "Merancang surat persetujuan", "", "", kotak, "", "", "", "Draft", "2 hari", "Draft"],
    ["4", "Telaah dan beri persetujuan", "", "", "", kotak, "", "", "Draft", "1 hari", "Persetujuan"],
    ["5", "Verifikasi + Tinjauan Lapangan", "", kotak, "", "", "", "", "Form + Bukti Lapangan", "3 hari", "Berita Acara"],
    ["6", "Cek kelengkapan", kotak, "", "", "", "", "", "📄", "15 menit", "✔️ atau ❌"],
    ["7", "Keputusan Izin", "", kotak, "", "", diamond, "", "Naskah Izin", "1 hari", "Surat Izin"]
]

columns = ["No", "Kegiatan", "Kasubbid Perizinan", "Kabid WASDAL", "Sekretaris", "Kaban", "Kabag Hukum", "Bupati", "Kelengkapan", "Waktu", "Output"]
df = pd.DataFrame(data, columns=columns)

# Tampilkan tabel di Streamlit
st.markdown(df.to_html(escape=False, index=False), unsafe_allow_html=True)

# Fungsi untuk membuat dokumen Word (Legal Landscape)
def generate_word_landscape(df):
    doc = Document()

    # Ubah ukuran halaman jadi Legal Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(35.56)
    section.page_height = Cm(21.59)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Judul
    doc.add_heading("📋 SOP - Penerbitan Izin Pemanfaatan Air Limbah", level=1)
    doc.add_paragraph("Bagian Proses dan Pelaksana (dengan simbol visual sederhana)")

    # Buat tabel
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    # Isi tabel
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    # Simpan ke buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Tombol Download
st.markdown("---")
st.subheader("⬇️ Unduh Dokumen Word - Legal Landscape")
word_file = generate_word_landscape(df)
st.download_button(
    label="📄 Download Word (Legal Landscape)",
    data=word_file,
    file_name="SOP_Izin_Landscape_Legal.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
