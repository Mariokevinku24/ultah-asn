import streamlit as st
from graphviz import Source
import tempfile
from PIL import Image
import base64
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Cm
from io import BytesIO
import os

st.set_page_config(layout="wide")
st.title("📊 Struktur Organisasi DLH Deli Serdang + Export Word")

# Sidebar Input Nama
st.sidebar.header("🖊️ Nama Pejabat")
kadis = st.sidebar.text_input("Kepala Dinas", "Dinas Lingkungan Hidup")
sekretaris = st.sidebar.text_input("Sekretaris", "SEKRETARIAT")
sub_umum = st.sidebar.text_input("Sub Bagian Umum", "Sub Bagian Umum")
sub_keuangan = st.sidebar.text_input("Sub Bagian Keuangan", "Sub Bagian Keuangan")
sub_program = st.sidebar.text_input("Sub Bagian Program", "Sub Bagian Program")

# Graphviz source
graph_code = f"""
digraph G {{
    rankdir=TB;
    node [shape=box, style="filled", fillcolor="#E0F7FA", fontname=Helvetica];

    "{kadis}";
    "{sekretaris}" -> "{sub_umum}";
    "{sekretaris}" -> "{sub_keuangan}";
    "{sekretaris}" -> "{sub_program}";
    "{kadis}" -> "{sekretaris}";
    "{kadis}" -> "Kelompok Jabatan\\nFungsional";

    "{kadis}" -> "Bidang Tata\\nLingkungan";
    "{kadis}" -> "Bidang Pengelolaan\\nSampah & Limbah B3";
    "{kadis}" -> "Bidang Pengendalian\\nPencemaran & Kerusakan";
    "{kadis}" -> "Bidang Penataan &\\nPeningkatan Kapasitas";

    "Bidang Tata\\nLingkungan" -> "Seksi Inventarisasi\\nRPPLH & KLHS";
    "Bidang Tata\\nLingkungan" -> "Seksi Kajian Dampak\\nLingkungan";
    "Bidang Tata\\nLingkungan" -> "Seksi Pemeliharaan\\nLingkungan Hidup";

    "Bidang Pengelolaan\\nSampah & Limbah B3" -> "Seksi Pengurangan\\nSampah";
    "Bidang Pengelolaan\\nSampah & Limbah B3" -> "Seksi Penanganan\\nSampah";
    "Bidang Pengelolaan\\nSampah & Limbah B3" -> "Seksi Limbah B3";

    "Bidang Pengendalian\\nPencemaran & Kerusakan" -> "Seksi Pemantauan\\nLingkungan";
    "Bidang Pengendalian\\nPencemaran & Kerusakan" -> "Seksi Pencemaran\\nLingkungan";
    "Bidang Pengendalian\\nPencemaran & Kerusakan" -> "Seksi Kerusakan\\nLingkungan";

    "Bidang Penataan &\\nPeningkatan Kapasitas" -> "Seksi Pengaduan &\\nPenyelesaian Sengketa";
    "Bidang Penataan &\\nPeningkatan Kapasitas" -> "Seksi Penegakan\\nHukum Lingkungan";
    "Bidang Penataan &\\nPeningkatan Kapasitas" -> "Seksi Peningkatan\\nKapasitas LH";

    "{kadis}" -> "UPT";
}}
"""

# Tampilkan grafik di Streamlit
st.subheader("📌 Struktur Organisasi")
st.graphviz_chart(graph_code)

# Buat Word dari grafik
if st.button("📥 Download Struktur dalam Word"):
    with st.spinner("Sedang memproses dokumen..."):
        # Render Graphviz jadi PNG
        src = Source(graph_code, format="png")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
            tmp_path = tmpfile.name
        src.render(filename=tmp_path, cleanup=True)
        img_path = tmp_path + ".png"

        # Buat dokumen Word
        doc = Document()

        # Atur kertas ke Legal Landscape
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_height = Cm(35.56)
        section.page_width = Cm(21.59)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

        doc.add_heading("Struktur Organisasi Dinas Lingkungan Hidup", level=1)
        doc.add_paragraph(f"Kepala Dinas: {kadis}")
        doc.add_paragraph(f"Sekretaris: {sekretaris}")

        # Tambahkan gambar
        doc.add_picture(img_path, width=Cm(26))  # sesuaikan agar cukup di landscape

        # Simpan ke buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📄 Download File Word",
            data=buffer,
            file_name="Struktur_Organisasi_DLH.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Cleanup sementara
        os.remove(img_path)

